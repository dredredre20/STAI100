"""
Cover letter generation: builds a tailored cover letter via the LLM and
compiles it into a DOCX file. Split out of react_agent.py so the
document-generation logic for cover letters lives independently of the
agent's tool-dispatch loop.
"""

import json
import os
import re
from datetime import datetime
from typing import Callable

from docx import Document
from docx.shared import Inches, Pt

from memory.db_memory import get_latest_resume_profile
from ds_integration.ingest_job_postings import job_collection
from ds_integration.job_fit_prediction import format_resume_string
from ds_integration.skill_gap import get_skill_gap_analysis


def generate_docx_cover_letter(cover_letter_text: str, output_path: str) -> str:
    """
    Compiles the provided cover letter plain text into a DOCX document.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = cover_letter_text.strip().split("\n")
    for line in lines:
        cleaned_line = line.strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if not cleaned_line:
            continue

        tokens = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
            else:
                p.add_run(token)

    doc.save(output_path)
    return output_path


def run_generate_cover_letter(
    params: dict,
    session_id: str,
    call_llm: Callable[[list], str],
) -> str:
    """
    Executes the `generate_cover_letter` tool end-to-end: fetches the user's
    resume profile, runs a gap analysis, asks the LLM to draft the letter,
    saves it as a DOCX, and returns a JSON string with the result.

    `call_llm` is injected by the caller (react_agent.py) to avoid a
    circular import between this module and the agent module.
    """
    try:
        company = params.get("company")
        title = params.get("title")
        if not company or not title:
            return json.dumps({"error": "generate_cover_letter requires both 'company' and 'title' parameters."})

        profile = get_latest_resume_profile(session_id)
        if not profile:
            return json.dumps({"note": "No resume profile found for this session to build a cover letter."})

        profile_dict = {
            "fields": {
                "current_role_category": profile.get("current_role_category"),
                "target_role": profile.get("target_role"),
                "years_of_experience": profile.get("years_of_experience"),
                "skills": json.loads(profile.get("skills") or "[]"),
                "certifications": json.loads(profile.get("certifications") or "[]"),
                "education_level": profile.get("education_level"),
            }
        }
        resume_text = format_resume_string(profile_dict)

        gap_analysis = {}
        try:
            gap_analysis = get_skill_gap_analysis(
                collection=job_collection,
                resume_text=resume_text,
                company=company,
                title=title,
            )
        except Exception:
            pass

        current_date = datetime.now().strftime("%B %d, %Y")

        prompt_messages = [
            {
                "role": "system",
                "content": (
                    "You are a professional resume writer. Write a formal, tailored cover letter "
                    "following standard professional formatting rules (Header, Salutation, Opening Hook, "
                    "Body Paragraph with matched achievements, Closing Call-to-Action, and Sign-off). "
                    "Output ONLY plain text without markdown code blocks."
                )
            },
            {
                "role": "user",
                "content": f"""
Target Company: {company}
Target Job Title: {title}

Candidate Background:
- Role/Category: {profile.get("current_role_category", "Professional")}
- Experience: {profile.get("years_of_experience", "N/A")} years
- Education: {profile.get("education_level", "N/A")}
- Technical Skills: {profile.get("skills", "[]")}
- Certifications: {profile.get("certifications", "[]")}

Job Requirements & Gap Analysis:
{json.dumps(gap_analysis, indent=2)}

Format Requirements:
1. HEADER:
   [Applicant Name]
   [City, State / Email / Phone Placeholder]
   {current_date}

   Hiring Manager or Hiring Team
   {company}

2. SALUTATION:
   Dear Hiring Team at {company},

3. OPENING PARAGRAPH:
   State job title ({title}), express enthusiasm for {company}, and outline top qualifications.

4. BODY PARAGRAPH(S):
   Detail relevant technical experience matching {company}'s requirements.

5. CLOSING PARAGRAPH:
   Reiterate enthusiasm and request an interview.

6. SIGN-OFF:
   Sincerely,
   [Applicant Name]
"""
            }
        ]

        cover_letter_text = call_llm(prompt_messages)

        clean_company = company.strip()
        docx_filename = f"Cover letter - {clean_company}.docx"
        docx_dir = os.path.join(os.getcwd(), "output", "cover_letters")
        docx_path = os.path.join(docx_dir, docx_filename)

        generate_docx_cover_letter(cover_letter_text, docx_path)

        return json.dumps({
            "company": company,
            "title": title,
            "cover_letter": cover_letter_text,
            "docx_path": docx_path
        })
    except Exception as e:
        return f"ERROR: generate_cover_letter failed: {e}"
