"""
Targeted resume generation: builds a Harvard-format tailored resume via the
LLM and compiles it into a DOCX file.

This module is designed to be called by the ReAct agent as part of the
resume-generation workflow. It:
1. Loads the user’s latest profile from memory,
2. Compares it against a target job via skill-gap analysis,
3. Prompts the LLM to draft a Harvard-formatted resume,
4. Saves the result as a DOCX file,
5. Returns the generated text and path as a JSON payload.
"""

import json
import os
import re
from typing import Callable

from docx import Document
from docx.shared import Inches, Pt

from memory.db_memory import get_latest_resume_profile
from ds_integration.ingest_job_postings import job_collection
from ds_integration.job_fit_prediction import format_resume_string
from ds_integration.skill_gap import get_skill_gap_analysis


def generate_docx_resume(resume_text: str, output_path: str) -> str:
    """
    Compile a plain-text resume into a DOCX file.

    The function converts the generated resume into a document while preserving
    bold formatting for section labels such as **EDUCATION** and **EXPERIENCE**.
    This keeps the output readable and visually close to a standard resume layout.

    Args:
        resume_text: Plain text resume content returned by the LLM.
        output_path: Full filesystem path where the DOCX file should be created.

    Returns:
        The saved DOCX path as a string.
    """

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = resume_text.strip().split("\n")
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15

        tokens = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
            else:
                p.add_run(token)

    doc.save(output_path)
    return output_path


def run_generate_targeted_resume(
    params: dict,
    session_id: str,
    call_llm: Callable[[list], str],
) -> str:
    """
    Generate a job-targeted resume end-to-end.

    This function:
    - reads the user profile from the active session,
    - retrieves a job/company context,
    - analyzes the skill gap between the profile and the target role,
    - prompts the LLM to generate a Harvard-formatted resume,
    - writes the document to disk as DOCX,
    - returns a JSON payload containing both the text and output path.

    Args:
        params: Request payload containing at least "company" and "title".
        session_id: Active conversation/session ID for loading stored profile data.
        call_llm: Injected LLM callback used to avoid circular imports.

    Returns:
        JSON string containing either the generated resume payload or an error note.
    """
    
    try:
        company = params.get("company")
        title = params.get("title")
        if not company or not title:
            return json.dumps({"error": "generate_targeted_resume requires both 'company' and 'title' parameters."})

        profile = get_latest_resume_profile(session_id)
        if not profile:
            return json.dumps({"note": "No resume profile found for this session."})

        profile_dict = {
            "fields": {
                "current_role_category": profile.get("current_role_category"),
                "target_role": profile.get("target_role"),
                "years_of_experience": profile.get("years_of_experience"),
                "skills": json.loads(profile.get("skills") or "[]"),
                "certifications": json.loads(profile.get("certifications") or "[]"),
                "education_level": profile.get("education_level"),
            }
        }
        original_resume_text = format_resume_string(profile_dict)

        gap_analysis = {}
        try:
            gap_analysis = get_skill_gap_analysis(
                collection=job_collection,
                resume_text=original_resume_text,
                company=company,
                title=title,
            )
        except Exception:
            pass

        prompt_messages = [
            {
                "role": "system",
                "content": (
                    "You are an expert executive resume writer. Generate a professional resume adhering strictly "
                    "to the official Harvard Resume Format guidelines. Tailor the content specifically for the target job "
                    "and company by combining original resume history with all recently updated skills."
                )
            },
            {
                "role": "user",
                "content": f"""
Target Company: {company}
Target Job Title: {title}

--- CANDIDATE DATA ---
Original Resume Content:
{original_resume_text}

Latest Profile Data (Includes skills added/updated during chat):
- Education Level: {profile.get("education_level", "N/A")}
- Years of Experience: {profile.get("years_of_experience", "N/A")}
- All Current & Updated Skills: {json.dumps(profile.get("skills", []))}
- Certifications: {json.dumps(profile.get("certifications", []))}

Job Gap Analysis:
{json.dumps(gap_analysis, indent=2)}

--- HARVARD RESUME FORMATTING & CONTENT RULES ---
Structure the resume into the following sections:

**FirstName LastName**
Address • City, State Zip • email@domain.com • Phone Number

**EDUCATION**
Institution Name | City, State or Country
Degree, Major / Concentration | Graduation Date
- Relevant Coursework or Academic Honors (if applicable)

**EXPERIENCE**
Organization Name | City, State or Country
Position Title | Month Year – Month Year
- Action-oriented bullet points outlining accomplishments, skills, and resulting outcomes.

**LEADERSHIP & ACTIVITIES**
Organization Name | City, State
Role | Month Year – Month Year
- Action-oriented bullet points highlighting leadership and initiative.

**SKILLS & INTERESTS**
- Technical: List software, frameworks, programming languages, and tools (must include all updated skills).
- Language: List foreign languages and level of fluency (if applicable).
- Interests: List activities or topics of interest that spark interview conversation.

MANDATORY BULLET RULES:
- Begin EVERY bullet line in Experience and Leadership with a strong action verb.
- Quantify accomplishments where possible (e.g., percentages, metrics, numbers).
- DO NOT use personal pronouns (I, me, my, we). Each line MUST be a phrase rather than a full sentence.
- Weave the candidate's updated skills into the Experience bullet points where relevant.
"""
            }
        ]

        targeted_resume_text = call_llm(prompt_messages)

        clean_company = company.strip()
        docx_filename = f"Targeted Resume - {clean_company}.docx"
        docx_dir = os.path.join(os.getcwd(), "output", "resumes")
        docx_path = os.path.join(docx_dir, docx_filename)

        generate_docx_resume(targeted_resume_text, docx_path)

        return json.dumps({
            "company": company,
            "title": title,
            "targeted_resume": targeted_resume_text,
            "docx_path": docx_path
        })
    except Exception as e:
        return f"ERROR: generate_targeted_resume failed: {e}"
