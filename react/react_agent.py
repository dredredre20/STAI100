import json
import os
import re
import ollama
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt

from config import MODEL, OLLAMA_BASE_URL
from memory.db_memory import (
    get_latest_resume_profile,
    update_skills
)
from memory.conversation_memory import (
    get_formatted_session_history,
    save_session_turn
)
from ds_integration.ingest_job_postings import job_collection
from ds_integration.job_fit_prediction import predict_job_matches, format_resume_string
from ds_integration.job_search import get_all_postings, search_postings
from ds_integration.skill_gap import get_skill_gap_analysis

TOOL_DESCRIPTIONS = """
Available Tools:
- get_top_matches[] : Ranks all stored job postings against the user's resume
  using semantic similarity, returning the top matches readiness tiers 
  (Ready to Apply / Near Match / Long-Term Upskilling).
  Use this when the user asks what jobs they're qualified for, what they
  should apply to, or wants an overview of their best-fit postings.

- get_skill_gap[company, title] : Retrieves a specific job posting (by company
  and/or title — at least one is required) and compares it against the user's
  resume. Returns matched skills and gaps, with the response style adapted to
  fit tier. Use this when the user names a specific company or role and asks what 
  they're missing for it, or asks for a detailed breakdown beyond a similarity score.

- search_posting[query] : Semantic search over stored job postings for a
  fuzzy/descriptive query (e.g. "AI engineering roles", "remote data roles").
  Returns matching postings with title/company/link and similarity. Use this
  when the user doesn't name an exact company/title but describes what
  they're looking for, or wants to browse rather than get a specific fit analysis.

- update_skills[new_skills, new_certifications] : Adds new skills and/or
  certifications to the user's saved resume profile.
  Use this when the user mentions earning a certification or learning a new skill.

- get_user_profile[] : Returns this session's most recently saved resume
  profile — skills, certifications, education level, years of experience,
  and target_role.

- generate_cover_letter[company, title] : Generates a customized cover letter and DOCX. 
  CRITICAL: Use this ONLY when the user explicitly asks to draft, write, or generate 
  a cover letter (e.g., "write a cover letter for X"). DO NOT trigger this for general 
  questions, advice, or skill gap queries. Both `company` and `title` parameters are required.

- generate_targeted_resume[company, title] : Generates a full, professionally formatted resume DOCX. 
  CRITICAL: Use this ONLY when the user explicitly asks to generate, build, or format 
  a resume (e.g., "generate a tailored resume for X"). DO NOT trigger this during standard 
  conversations or skill gap queries. Both `company` and `title` parameters are required.
"""

SYSTEM_PROMPT_TEMPLATE = """Role: You are a career-readiness advisor agent.
You help users understand their skill gaps for a target role and track their progress.
{tool_descriptions}
User context:
- session_id: {session_id}
- resume_skills: {resume_skills}
- target_role: {target_role}

CONVERSATION HISTORY (rolling summary + recent turns from earlier in this session):
{conversation_history}
---

RULES FOR TOOL CALLING:
1. NEVER call `generate_cover_letter` or `generate_targeted_resume` unless the user explicitly requested document creation in their latest prompt.
2. If the user asks general questions about job fit, skill gaps, missing skills, or career advice, use `get_skill_gap` or `get_top_matches` instead.

You must respond with ONLY a valid JSON object. Do not wrap it in markdown code blocks.
Do not include any explanatory text before or after the JSON.

Every response has this shape:
{{"thought": "...", "action": {{"tool_name": "...", "parameters": {{...}}}} or null, "final_answer": "..." or null}}

Exactly one of "action" or "final_answer" must be non-null. The other must be null.

Conventions:
- When calling `generate_cover_letter` or `generate_targeted_resume`, your `final_answer` in the following turn MUST output the exact text of the document AND notify the user that the DOCX file has been generated and saved by stating exactly: "[DOCX Generated]: Saved to <path>".
- Be concise and conversational in final_answer — this is a chat interface, not a report.
"""


def _get_resume_profile_dict(session_id: str) -> dict:
    profile = get_latest_resume_profile(session_id)
    if not profile:
        return {}
    return {
        "fields": {
            "current_role_category": profile.get("current_role_category"),
            "target_role": profile.get("target_role"),
            "years_of_experience": profile.get("years_of_experience"),
            "skills": json.loads(profile.get("skills") or "[]"),
            "certifications": json.loads(profile.get("certifications") or "[]"),
            "education_level": profile.get("education_level"),
        }
    }


def generate_docx_cover_letter(cover_letter_text: str, output_path: str) -> str:
    """
    Compiles the provided cover letter plain text into a DOCX document.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = cover_letter_text.strip().split("\n")
    for line in lines:
        cleaned_line = line.strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if not cleaned_line:
            continue

        tokens = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
            else:
                p.add_run(token)

    doc.save(output_path)
    return output_path


def generate_docx_resume(resume_text: str, output_path: str) -> str:
    """
    Compiles the tailored resume text into a Harvard-formatted DOCX document safely.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = resume_text.strip().split("\n")
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15

        tokens = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
            else:
                p.add_run(token)

    doc.save(output_path)
    return output_path


def run_tool(
    action: dict,
    session_id: str,
    resume_skills: list[str],
    target_role: str,
    user_message: str = ""
) -> str:
    tool_name = action.get("tool_name", "")
    params = action.get("parameters", {})

    # Code-level guardrail against over-triggering document tools
    if tool_name in ["generate_cover_letter", "generate_targeted_resume"]:
        trigger_words = ["cover letter", "resume", "write", "generate", "draft", "create", "build", "make"]
        if user_message and not any(word in user_message.lower() for word in trigger_words):
            return json.dumps({
                "error": f"Tool '{tool_name}' was blocked because the user did not explicitly request document creation. Answer their query conversationally instead."
            })

    if tool_name == "get_user_profile":
        try:
            profile = get_latest_resume_profile(session_id)
            if not profile:
                return json.dumps({"note": "No resume profile found for this session."})
            return json.dumps({
                "target_role": profile.get("target_role"),
                "current_role_category": profile.get("current_role_category"),
                "years_of_experience": profile.get("years_of_experience"),
                "skills": json.loads(profile.get("skills") or "[]"),
                "certifications": json.loads(profile.get("certifications") or "[]"),
                "education_level": profile.get("education_level"),
            })
        except Exception as e:
            return f"ERROR: get_user_profile failed: {e}"

    elif tool_name == "update_skills":
        try:
            new_skills = params.get("new_skills", [])
            new_certifications = params.get("new_certifications", [])
            if not new_skills and not new_certifications:
                return json.dumps({"error": "update_skills requires at least one new skill or certification."})

            result = update_skills(
                session_id=session_id,
                new_skills=new_skills,
                new_certifications=new_certifications,
            )
            return json.dumps(result)
        except Exception as e:
            return f"ERROR: update_skills failed: {e}"

    elif tool_name == "get_top_matches":
        try:
            profile_dict = _get_resume_profile_dict(session_id)
            if not profile_dict:
                return json.dumps({"note": "No resume profile found for this session."})
            resume_text = format_resume_string(profile_dict)

            all_postings = get_all_postings()
            if not all_postings:
                return json.dumps({"note": "No job postings are currently stored."})

            results_df = predict_job_matches(resume_text, all_postings)
            top_matches = results_df.head(5)[["title", "company", "link", "tier"]].to_dict(orient="records")
            return json.dumps({"top_matches": top_matches})
        except Exception as e:
            return f"ERROR: get_top_matches failed: {e}"

    elif tool_name == "get_skill_gap":
        try:
            company = params.get("company")
            title = params.get("title")
            if not company and not title:
                return json.dumps({"error": "get_skill_gap requires at least a company or title."})

            profile_dict = _get_resume_profile_dict(session_id)
            if not profile_dict:
                return json.dumps({"note": "No resume profile found for this session."})
            resume_text = format_resume_string(profile_dict)

            result = get_skill_gap_analysis(
                collection=job_collection,
                resume_text=resume_text,
                company=company,
                title=title,
            )
            return json.dumps(result)
        except Exception as e:
            return f"ERROR: get_skill_gap failed: {e}"

    elif tool_name == "search_posting":
        try:
            query = params.get("query", "")
            if not query:
                return json.dumps({"error": "search_posting requires a query."})
            matches = search_postings(query, n_results=5)
            return json.dumps({"results": matches})
        except Exception as e:
            return f"ERROR: search_posting failed: {e}"

    elif tool_name == "generate_cover_letter":
        try:
            company = params.get("company")
            title = params.get("title")
            if not company or not title:
                return json.dumps({"error": "generate_cover_letter requires both 'company' and 'title' parameters."})

            profile = get_latest_resume_profile(session_id)
            if not profile:
                return json.dumps({"note": "No resume profile found for this session to build a cover letter."})

            profile_dict = _get_resume_profile_dict(session_id)
            resume_text = format_resume_string(profile_dict)

            gap_analysis = {}
            try:
                gap_analysis = get_skill_gap_analysis(
                    collection=job_collection,
                    resume_text=resume_text,
                    company=company,
                    title=title,
                )
            except Exception:
                pass

            current_date = datetime.now().strftime("%B %d, %Y")

            prompt_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are a professional resume writer. Write a formal, tailored cover letter "
                        "following standard professional formatting rules (Header, Salutation, Opening Hook, "
                        "Body Paragraph with matched achievements, Closing Call-to-Action, and Sign-off). "
                        "Output ONLY plain text without markdown code blocks."
                    )
                },
                {
                    "role": "user",
                    "content": f"""
Target Company: {company}
Target Job Title: {title}

Candidate Background:
- Role/Category: {profile.get("current_role_category", "Professional")}
- Experience: {profile.get("years_of_experience", "N/A")} years
- Education: {profile.get("education_level", "N/A")}
- Technical Skills: {profile.get("skills", "[]")}
- Certifications: {profile.get("certifications", "[]")}

Job Requirements & Gap Analysis:
{json.dumps(gap_analysis, indent=2)}

Format Requirements:
1. HEADER:
   [Applicant Name]
   [City, State / Email / Phone Placeholder]
   {current_date}

   Hiring Manager or Hiring Team
   {company}

2. SALUTATION:
   Dear Hiring Team at {company},

3. OPENING PARAGRAPH:
   State job title ({title}), express enthusiasm for {company}, and outline top qualifications.

4. BODY PARAGRAPH(S):
   Detail relevant technical experience matching {company}'s requirements.

5. CLOSING PARAGRAPH:
   Reiterate enthusiasm and request an interview.

6. SIGN-OFF:
   Sincerely,
   [Applicant Name]
"""
                }
            ]

            cover_letter_text = call_llm(prompt_messages)

            clean_company = company.strip()
            docx_filename = f"Cover letter - {clean_company}.docx"
            docx_dir = os.path.join(os.getcwd(), "output", "cover_letters")
            docx_path = os.path.join(docx_dir, docx_filename)
            
            generate_docx_cover_letter(cover_letter_text, docx_path)

            return json.dumps({
                "company": company,
                "title": title,
                "cover_letter": cover_letter_text,
                "docx_path": docx_path
            })
        except Exception as e:
            return f"ERROR: generate_cover_letter failed: {e}"

    elif tool_name == "generate_targeted_resume":
        try:
            company = params.get("company")
            title = params.get("title")
            if not company or not title:
                return json.dumps({"error": "generate_targeted_resume requires both 'company' and 'title' parameters."})

            profile = get_latest_resume_profile(session_id)
            if not profile:
                return json.dumps({"note": "No resume profile found for this session."})

            profile_dict = _get_resume_profile_dict(session_id)
            original_resume_text = format_resume_string(profile_dict)

            gap_analysis = {}
            try:
                gap_analysis = get_skill_gap_analysis(
                    collection=job_collection,
                    resume_text=original_resume_text,
                    company=company,
                    title=title,
                )
            except Exception:
                pass

            prompt_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an expert executive resume writer. Generate a professional resume adhering strictly "
                        "to the official Harvard Resume Format guidelines. Tailor the content specifically for the target job "
                        "and company by combining original resume history with all recently updated skills."
                    )
                },
                {
                    "role": "user",
                    "content": f"""
Target Company: {company}
Target Job Title: {title}

--- CANDIDATE DATA ---
Original Resume Content:
{original_resume_text}

Latest Profile Data (Includes skills added/updated during chat):
- Education Level: {profile.get("education_level", "N/A")}
- Years of Experience: {profile.get("years_of_experience", "N/A")}
- All Current & Updated Skills: {json.dumps(profile.get("skills", []))}
- Certifications: {json.dumps(profile.get("certifications", []))}

Job Gap Analysis:
{json.dumps(gap_analysis, indent=2)}

--- HARVARD RESUME FORMATTING & CONTENT RULES ---
Structure the resume into the following sections:

**FirstName LastName**
Address • City, State Zip • email@domain.com • Phone Number

**EDUCATION**
Institution Name | City, State or Country
Degree, Major / Concentration | Graduation Date
- Relevant Coursework or Academic Honors (if applicable)

**EXPERIENCE**
Organization Name | City, State or Country
Position Title | Month Year – Month Year
- Action-oriented bullet points outlining accomplishments, skills, and resulting outcomes.

**LEADERSHIP & ACTIVITIES**
Organization Name | City, State
Role | Month Year – Month Year
- Action-oriented bullet points highlighting leadership and initiative.

**SKILLS & INTERESTS**
- Technical: List software, frameworks, programming languages, and tools (must include all updated skills).
- Language: List foreign languages and level of fluency (if applicable).
- Interests: List activities or topics of interest that spark interview conversation.

MANDATORY BULLET RULES:
- Begin EVERY bullet line in Experience and Leadership with a strong action verb.
- Quantify accomplishments where possible (e.g., percentages, metrics, numbers).
- DO NOT use personal pronouns (I, me, my, we). Each line MUST be a phrase rather than a full sentence.
- Weave the candidate's updated skills into the Experience bullet points where relevant.
"""
                }
            ]

            targeted_resume_text = call_llm(prompt_messages)

            clean_company = company.strip()
            docx_filename = f"Targeted Resume - {clean_company}.docx"
            docx_dir = os.path.join(os.getcwd(), "output", "resumes")
            docx_path = os.path.join(docx_dir, docx_filename)
            
            generate_docx_resume(targeted_resume_text, docx_path)

            return json.dumps({
                "company": company,
                "title": title,
                "targeted_resume": targeted_resume_text,
                "docx_path": docx_path
            })
        except Exception as e:
            return f"ERROR: generate_targeted_resume failed: {e}"
    else:
        return f"ERROR: Unknown tool '{tool_name}'"


def format_final_answer(answer: str) -> str:
    try:
        data = json.loads(answer)
    except (json.JSONDecodeError, TypeError):
        return answer

    if not isinstance(data, dict):
        return answer

    parts = []
    if data.get("cover_letter"):
        parts.append(data["cover_letter"])
        if data.get("docx_path"):
            parts.append(f"\n\n[DOCX Generated]: Saved to {data['docx_path']}")
            
    if data.get("targeted_resume"):
        parts.append(data["targeted_resume"])
        if data.get("docx_path"):
            parts.append(f"\n\n[DOCX Generated]: Saved to {data['docx_path']}")
            
    if data.get("top_matches"):
        titles = [f"{m.get('title')} at {m.get('company')} ({m.get('tier')})" for m in data["top_matches"]]
        parts.append("Top matches: " + ", ".join(titles) + ".")
    if data.get("results"):
        titles = [f"{r.get('title')} at {r.get('company')}" for r in data["results"]]
        parts.append("Found: " + ", ".join(titles) + ".")
    if data.get("summary"):
        parts.append(data["summary"])
    return " ".join(parts) if parts else answer


def run_agent(
    user_message: str,
    session_id: str,
    resume_skills: list[str],
    target_role: str,
    max_turns: int = 10,
    verbose: bool = True,
) -> str:

    session_history, conversation_history_text = get_formatted_session_history(session_id)

    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(
        tool_descriptions=TOOL_DESCRIPTIONS,
        session_id=session_id,
        resume_skills=json.dumps(resume_skills),
        target_role=target_role,
        conversation_history=conversation_history_text,
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
    ]

    final_answer = None
    for turn in range(1, max_turns + 1):
        if verbose:
            print(f"\n------ TURN {turn} ------")
        try:
            raw = call_llm(messages)
            data = parse_json(raw)
        except Exception as e:
            if verbose:
                print(f"[Error] {e}")
            final_answer = "Sorry, I ran into an error trying to answer that."
            break

        messages.append({"role": "assistant", "content": raw})
        if verbose:
            print(f"Thought: {data.get('thought', '')}")

        action = data.get("action")
        if action:
            if verbose:
                print(f"Action: {json.dumps(action)}")
            result_str = run_tool(
                action=action,
                session_id=session_id,
                resume_skills=resume_skills,
                target_role=target_role,
                user_message=user_message
            )
            if verbose:
                print(f"Result: {result_str[:300]}")
            messages.append({"role": "user", "content": f"Observation: {result_str}"})
        else:
            final_answer = format_final_answer(data.get("final_answer", ""))
            break

    if final_answer is None:
        final_answer = "I wasn't able to reach an answer within the allowed number of steps."

    save_session_turn(session_id, session_history, user_message, final_answer)

    return final_answer


def call_llm(messages: list) -> str:
    client = ollama.Client(host=OLLAMA_BASE_URL)
    response = client.chat(model=MODEL, messages=messages, format="json" if len(messages) > 2 else None)
    return response['message']['content'].strip()


def parse_json(text: str) -> dict:
    text = text.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    return json.loads(text)